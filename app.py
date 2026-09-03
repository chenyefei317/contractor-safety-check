import datetime
import io
import zipfile
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt
import numpy as np
from PIL import Image
import pandas as pd
import streamlit as st
from streamlit_drawable_canvas import st_canvas

st.set_page_config(page_title="承包商入场EHS审核与承诺书", layout="centered")

# === 隐藏右上角 Streamlit 默认菜单、页脚及源码查看通道 ===
hide_streamlit_style = """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    </style>
    """
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# 侧边栏：手机端/网页端扫码填报支持
with st.sidebar:
  st.header("📱 扫码/网页端填报")
  st.write("承包商可通过手机扫描二维码或直接打开链接进行填报与资料上传。")

  app_url = st.text_input(
      "本应用公网链接 (URL):",
      value="https://contractor-safety-check.streamlit.app/",
  )

  if app_url:
    qr_api_url = f"https://api.qrserver.com/v1/create-qr-code/?size=200x200&data={app_url}"
    st.image(qr_api_url, caption="手机相机/支付宝/浏览器扫码", width=200)

# === 顶部 Logo (尺寸调整为 width=120) ===
try:
  st.image("logo.png", width=120)
except Exception:
  pass

st.title("👷 承包商入场EHS审核与承诺书")
st.write(
    "依据《安全生产法》及 **宜家供应商IWAY6.0合规风险与管控要求**，请在入场前逐项核实、上传证明并签字。个人开发工具，严禁商业用途。"
)

# 1. 前置检查项 (已取消原第3项保险证明)
st.subheader("1. 承包商资信与资质审核 (入场必选)")
q1_1 = st.checkbox("1）企业资质类资料--基础证照：有效的营业执照（须上传）")
q1_2 = st.checkbox("2）提交甲乙双方盖章的承包商安全协议（须上传）")

# 2. 现场人员资质与基础保障
st.subheader("2. 现场人员资质与基础保障 (入场必选)")
q2_1 = st.checkbox(
    "1）提供承包商员工入厂安全培训记录（注：需在入场施工前完成并确认）。"
)
q2_2 = st.checkbox(
    "2）提供承包商施工人员工伤保险缴纳证明或意外伤害保险缴纳证明（须上传）。"
)
q2_3 = st.checkbox("3）作业人员的身份证复印件（须上传）。")

st.subheader("3. 特种/高危作业特批 (按需填写，无则留空)")
q3_1 = st.checkbox("【动火作业】如涉及动火或切割产生明火作业，申请报备。")
id_3_1 = st.text_input("备选特种作业人员 (动火) 身份证号：", key="id_fire")

q3_2 = st.checkbox("【电工作业】如涉及电工 (低压作业), 申请报备。")
id_3_2 = st.text_input("备选特种作业人员 (电工) 身份证号：", key="id_elec")

q3_3 = st.checkbox("【登高作业】如涉及登高作业，申请报备。")
id_3_3 = st.text_input("备选特种作业人员 (登高) 身份证号：", key="id_high")

# 4. 文件/证件上传区域
st.subheader("4. 相关证明材料上传")
st.write(
    "请上传营业执照、工伤保险证明、特种作业操作证等相关证件照片或扫描件（支持多选）。"
)
uploaded_files = st.file_uploader(
    "上传凭证/证件文件", type=["png", "jpg", "jpeg", "pdf"], accept_multiple_files=True
)

st.subheader("5. 自我承诺与声明")
st.info(
    "本人郑重承诺：以上勾选的核验项及上传的资料均真实有效，绝无弄虚作假。若因未落实安全防范措施而导致相关事故，愿承担相应的管理责任。"
)

declaration = st.checkbox("我已阅读并同意以上自我承诺")

col1, col2 = st.columns(2)
with col1:
  checker_name = st.text_input("承诺人姓名 (必填)：")
with col2:
  commit_date = st.date_input("承诺日期：", datetime.date.today())

st.write("---")
st.write("✍️ **请在下方空白处手写签名：**")
canvas_result = st_canvas(
    stroke_width=3,
    stroke_color="#000000",
    background_color="#F0F2F6",
    height=150,
    width=400,
    drawing_mode="freedraw",
    key="canvas",
)

# 提交并生成最终统一报告
if st.button("📁 确认无误，生成完整校验报告并打包下载"):
  base_passed = q1_1 and q1_2 and q2_1 and q2_2 and q2_3

  missing_ids = []
  if q3_1 and not id_3_1.strip():
    missing_ids.append("动火作业")
  if q3_2 and not id_3_2.strip():
    missing_ids.append("电工作业")
  if q3_3 and not id_3_3.strip():
    missing_ids.append("登高作业")

  if not base_passed:
    st.error(
        "❌ 警告：第 1 和第 2 部分为基础必选项！未全部落实前，绝对禁止办理入场。"
    )
  elif missing_ids:
    st.warning(
        f"⚠️ 拦截：您勾选了特种作业申请（{'、'.join(missing_ids)}），请务必在上方填写对应的作业人员身份证号！"
    )
  elif not declaration or not checker_name:
    st.warning("⚠️ 流程未完成：请勾选【自我承诺】并填写【承诺人姓名】。")
  elif canvas_result.image_data is None:
    st.warning("⚠️ 请在上方画板完成手写签名后再提交。")
  else:
    st.success("✅ 核验与承诺通过！以下为本次入场的完整合规档案：")

    status_3_1 = f"已报备 (身份证: {id_3_1})" if q3_1 else "未涉及此作业"
    status_3_2 = f"已报备 (身份证: {id_3_2})" if q3_2 else "未涉及此作业"
    status_3_3 = f"已报备 (身份证: {id_3_3})" if q3_3 else "未涉及此作业"

    data = {
        "安全核验管控项目": [
            "企业资质 - 有效的营业执照（须上传）",
            "企业资质 - 承包商安全协议（须上传）",
            "人员资质 - 入厂安全培训记录",
            "人员资质 - 施工人员保险证明（须上传）",
            "人员资质 - 身份证复印件（须上传）",
            "特种作业报备：动火作业",
            "特种作业报备：电工作业",
            "特种作业报备：登高作业",
        ],
        "现场确认结果": [
            "合格 / 承诺已上传",
            "合格 / 承诺已上传",
            "合格 / 需现场确认",
            "合格 / 承诺已上传",
            "合格 / 承诺已上传",
            status_3_1,
            status_3_2,
            status_3_3,
        ],
    }

    df = pd.DataFrame(data)

    # === 页面展示区 ===
    st.markdown("### 📋 承包商入场安全核验清单汇总")
    st.table(df)

    if uploaded_files:
      st.markdown(f"**📎 已上传证明文件数量：** {len(uploaded_files)} 个")

    st.markdown("---")
    st.markdown(f"**📝 声明承诺人：** {checker_name}")
    st.markdown(f"**📅 承诺生效日期：** {commit_date}")
    st.markdown("**✍️ 现场手写签名确认：**")

    signature_img = Image.fromarray(
        canvas_result.image_data.astype("uint8"), "RGBA"
    )
    st.image(signature_img, width=400, caption=f"承诺人：{checker_name}")

    # === 打印 / 另存为 PDF 按钮 ===
    st.markdown("---")
    st.markdown("### 🖨️ 文档操作选项")
    st.markdown(
        """
            <button onclick="window.print()" style="background-color: #ff4b4b; color: white; padding: 10px 20px; border: none; border-radius: 5px; cursor: pointer; font-weight: bold; font-size: 16px;">
                🖨️ 直接打印网页 / 另存为 PDF
            </button>
            """,
        unsafe_allow_html=True,
    )
    st.write(
        "*(点击上方按钮可直接调起浏览器的打印机或选择“另存为 PDF”)*"
    )

    # === 导出 Word (.docx) 文件流（华文宋体、取消所有粗体、仅“须上传”文字设为绿色且不加粗） ===
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "华文宋体"
    style.font.element.rPr.rFonts.set(qn("w:eastAsia"), "华文宋体")

    # 大标题：不加粗
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("承包商入场EHS审核与承诺书")
    run_h1.bold = False
    run_h1.font.name = "华文宋体"
    run_h1.font.element.rPr.rFonts.set(qn("w:eastAsia"), "华文宋体")

    p_info = doc.add_paragraph()
    run_info = p_info.add_run(
        f"承诺人：{checker_name}    承诺日期：{commit_date}"
    )
    run_info.bold = False
    run_info.font.name = "华文宋体"
    run_info.font.element.rPr.rFonts.set(qn("w:eastAsia"), "华文宋体")

    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run("安全核验管控项目清单：")
    run_sub.bold = False
    run_sub.font.name = "华文宋体"
    run_sub.font.element.rPr.rFonts.set(qn("w:eastAsia"), "华文宋体")

    table = doc.add_table(rows=len(df) + 1, cols=2)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "安全核验管控项目"
    hdr_cells[1].text = "现场确认结果"
    for cell in hdr_cells:
      for paragraph in cell.paragraphs:
        for run in paragraph.runs:
          run.bold = False
          run.font.name = "华文宋体"
          run.font.element.rPr.rFonts.set(qn("w:eastAsia"), "华文宋体")

    for i, row in df.iterrows():
      row_cells = table.rows[i + 1].cells
      proj_text = str(row["安全核验管控项目"])
      res_text = str(row["现场确认结果"])

      row_cells[0].text = ""
      p0 = row_cells[0].paragraphs[0]

      # 如果包含“须上传”，将“须上传”文字设为绿色（不加粗）
      if "须上传" in proj_text:
        parts = proj_text.split("（须上传）")
        run0_1 = p0.add_run(parts[0])
        run0_1.bold = False
        run0_1.font.name = "华文宋体"
        run0_1.font.element.rPr.rFonts.set(qn("w:eastAsia"), "华文宋体")

        run0_2 = p0.add_run("（须上传）")
        run0_2.bold = False
        run0_2.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
        run0_2.font.name = "华文宋体"
        run0_2.font.element.rPr.rFonts.set(qn("w:eastAsia"), "华文宋体")

        if len(parts) > 1 and parts[1]:
          run0_3 = p0.add_run(parts[1])
          run0_3.bold = False
          run0_3.font.name = "华文宋体"
          run0_3.font.element.rPr.rFonts.set(qn("w:eastAsia"), "华文宋体")
      else:
        run0 = p0.add_run(proj_text)
        run0.bold = False
        run0.font.name = "华文宋体"
        run0.font.element.rPr.rFonts.set(qn("w:eastAsia"), "华文宋体")

      row_cells[1].text = ""
      p1 = row_cells[1].paragraphs[0]
      run1 = p1.add_run(res_text)
      run1.bold = False
      run1.font.name = "华文宋体"
      run1.font.element.rPr.rFonts.set(qn("w:eastAsia"), "华文宋体")

    p_sig = doc.add_paragraph()
    run_sig = p_sig.add_run(f"\n声明承诺人签名：{checker_name}")
    run_sig.bold = False
    run_sig.font.name = "华文宋体"
    run_sig.font.element.rPr.rFonts.set(qn("w:eastAsia"), "华文宋体")

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    col_w, col_z = st.columns(2)
    with col_w:
      docx_filename = f"承包商入场核验_{checker_name}_{commit_date}.docx"
      st.download_button(
          label="📄 导出 Word 文档 (.docx)",
          data=doc_buffer,
          file_name=docx_filename,
          mime=(
              "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
          ),
      )

    # === 打包下载 ZIP ===
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
      csv_bytes = df.to_csv(index=False).encode("utf-8-sig")
      csv_filename = f"承包商入场核验_{checker_name}_{commit_date}.csv"
      zip_file.writestr(csv_filename, csv_bytes)

      img_byte_arr = io.BytesIO()
      signature_img.save(img_byte_arr, format="PNG")
      img_filename = f"签名_{checker_name}_{commit_date}.png"
      zip_file.writestr(img_filename, img_byte_arr.getvalue())

      if uploaded_files:
        for idx, uploaded_file in enumerate(uploaded_files):
          file_ext = uploaded_file.name.split(".")[-1]
          safe_file_name = (
              f"证明文件_{idx+1}_{checker_name}_{commit_date}.{file_ext}"
          )
          zip_file.writestr(safe_file_name, uploaded_file.getvalue())

    zip_buffer.seek(0)

    with col_z:
      zip_filename = f"承包商安全合规档案及证件_{checker_name}_{commit_date}.zip"
      st.download_button(
          label="📥 打包下载全部档案 (含清单、签名与凭证 ZIP)",
          data=zip_buffer,
          file_name=zip_filename,
          mime="application/zip",
      )

    st.balloons()
